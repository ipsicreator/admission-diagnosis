import io
import json
import re
import sqlite3
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple

import pandas as pd
import streamlit as st
from docx import Document
from pypdf import PdfReader


BASE_DIR = Path(__file__).resolve().parent
DATA_DIR = BASE_DIR / "data"
OUTPUT_DIR = BASE_DIR / "output"
DB_PATH = BASE_DIR / "admission_diag.db"

SUSI_2026 = DATA_DIR / "susi_explorer.csv"
SUSI_2027 = DATA_DIR / "susi_explorer_2027.csv"
CUTOFFS = DATA_DIR / "admission_cutoffs.csv"
CRITERIA = DATA_DIR / "holistic_criteria.csv"

TOP15 = [
    "서울대",
    "연세대",
    "고려대",
    "서강대",
    "성균관대",
    "한양대",
    "중앙대",
    "경희대",
    "한국외대",
    "서울시립대",
    "이화여대",
    "건국대",
    "동국대",
    "홍익대",
    "숙명여대",
]

GRADE_OPTIONS = ["고1", "고2", "고3", "N수이상"]
ADMISSION_TYPES = ["학생부교과", "학생부종합", "논술", "특기자/실기", "기타"]


@dataclass
class SupportChoice:
    support_no: int
    university: str
    department: str
    admission_type: str
    track_name: str
    diag_level: str
    diag_reason: str
    cutoff50: float
    cutoff70: float
    cutoff_basis: str


def inject_css() -> None:
    st.markdown(
        """
        <style>
        .service-title {
          text-align: center;
          font-size: 33px; /* 25pt */
          font-weight: 800;
          margin-bottom: 2px;
        }
        .service-subtitle {
          text-align: center;
          font-size: 20px; /* 15pt */
          font-weight: 600;
          margin-bottom: 14px;
          opacity: 0.95;
        }
        .step-box {
          text-align: center;
          font-size: 31px;
          font-weight: 700;
          padding: 12px 4px;
          border-radius: 10px;
        }
        .step-active {
          background: rgba(30, 91, 255, 0.18);
          border: 1px solid rgba(84, 139, 255, 0.55);
        }
        .footer-note {
          text-align: center;
          margin-top: 22px;
          opacity: 0.85;
          font-size: 14px;
        }
        .top-meta-left {
          font-size: 13px;
          opacity: 0.9;
          text-align: left;
          margin-bottom: 4px;
        }
        .top-meta-right {
          font-size: 13px;
          opacity: 0.9;
          text-align: right;
          margin-bottom: 4px;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )


def ensure_dirs() -> None:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def init_db() -> None:
    ensure_dirs()
    conn = sqlite3.connect(DB_PATH)
    try:
        cur = conn.cursor()
        cur.execute(
            """
            CREATE TABLE IF NOT EXISTS diagnosis_sessions (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                created_at TEXT NOT NULL,
                consultant_name TEXT,
                student_name TEXT,
                school_name TEXT,
                grade TEXT,
                student_phone TEXT,
                email TEXT,
                parent_phone TEXT,
                student_index REAL,
                pdf_summary TEXT,
                holistic_score REAL,
                outcome_json TEXT
            )
            """
        )
        cur.execute(
            """
            CREATE TABLE IF NOT EXISTS support_choices (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                session_id INTEGER NOT NULL,
                support_no INTEGER NOT NULL,
                university TEXT,
                department TEXT,
                admission_type TEXT,
                track_name TEXT,
                diag_level TEXT,
                diag_reason TEXT,
                cutoff50 REAL,
                cutoff70 REAL,
                FOREIGN KEY(session_id) REFERENCES diagnosis_sessions(id)
            )
            """
        )
        conn.commit()
    finally:
        conn.close()


def _normalize_columns(df: pd.DataFrame, kind: str) -> pd.DataFrame:
    alias_map = {
        "year": ["year", "연도", "학년도"],
        "university": ["university", "대학교", "대학명", "대학"],
        "department": ["department", "모집단위", "학과", "학부"],
        "admission_type": ["admission_type", "전형유형", "전형유형명"],
        "track_name": ["track_name", "전형명"],
        "percentile_type": ["percentile_type", "컷구분", "percentile"],
        "cutoff_score": ["cutoff_score", "컷점수", "점수", "score"],
    }
    rename_map: Dict[str, str] = {}
    cols = set(df.columns)
    for target, aliases in alias_map.items():
        for alias in aliases:
            if alias in cols:
                rename_map[alias] = target
                break
    df = df.rename(columns=rename_map)

    required = ["year", "university", "department", "admission_type", "track_name"]
    if kind == "cutoff":
        required += ["percentile_type", "cutoff_score"]

    for c in required:
        if c not in df.columns:
            df[c] = None

    return df[required].copy()


def load_csv(path: Path) -> pd.DataFrame:
    return pd.read_csv(path, encoding="utf-8-sig")


def save_uploaded_csv(uploaded_file, target_path: Path) -> None:
    target_path.write_bytes(uploaded_file.getvalue())


def _to_int_year(series: pd.Series) -> pd.Series:
    return pd.to_numeric(series, errors="coerce").fillna(-1).astype(int)


def _clean_text(s: pd.Series) -> pd.Series:
    return s.fillna("").astype(str).str.strip()


def _univ_key(name: str) -> str:
    s = str(name or "").strip().replace(" ", "")
    for suffix in ["대학교", "대학", "대"]:
        if s.endswith(suffix):
            s = s[: -len(suffix)]
    return s


def _canonical_university_map(base_2026: pd.DataFrame) -> Dict[str, str]:
    names = sorted(base_2026["university"].dropna().astype(str).str.strip().unique().tolist())
    key_map: Dict[str, str] = {}
    for n in names:
        k = _univ_key(n)
        if not k:
            continue
        if k not in key_map:
            key_map[k] = n
        else:
            prev = key_map[k]
            # 2026 명칭 중 더 공식형(길이 긴 값) 우선
            key_map[k] = n if len(n) > len(prev) else prev
    return key_map


def _remove_excluded_type(admission_type: str, track_name: str) -> bool:
    joined = (str(admission_type) + " " + str(track_name)).replace(" ", "")
    return "교과기회" in joined


def _extract_year_from_col(col_name: str) -> Optional[int]:
    text = str(col_name)
    for y in [2023, 2024, 2025]:
        if str(y) in text:
            return y
    return None


def _detect_percentile_from_text(text: str) -> Optional[int]:
    t = str(text).replace(" ", "")
    if "50%" in t or "50컷" in t:
        return 50
    if "70%" in t or "70컷" in t:
        return 70
    return None


def _normalize_cutoff_flexible(df: pd.DataFrame) -> pd.DataFrame:
    # long format
    if {"year", "university", "department", "admission_type", "track_name", "percentile_type", "cutoff_score"}.issubset(set(df.columns)):
        out = df.copy()
        out["year"] = _to_int_year(out["year"])
        out["percentile_type"] = pd.to_numeric(out["percentile_type"], errors="coerce")
        out["cutoff_score"] = pd.to_numeric(out["cutoff_score"], errors="coerce")
        for c in ["university", "department", "admission_type", "track_name"]:
            out[c] = _clean_text(out[c])
        return out[
            ["year", "university", "department", "admission_type", "track_name", "percentile_type", "cutoff_score"]
        ].copy()

    # flexible wide format (e.g. 2025학년도 입결(등급), 2024..., 2023...)
    rename_candidates = {
        "university": ["university", "대학교", "대학명", "대학"],
        "department": ["department", "모집단위", "학과", "학부"],
        "admission_type": ["admission_type", "전형유형", "전형유형명"],
        "track_name": ["track_name", "전형명"],
        "basis": ["기준", "2025학년도 기준", "컷기준", "구분"],
    }
    col_map: Dict[str, str] = {}
    cols = set(df.columns)
    for target, aliases in rename_candidates.items():
        for a in aliases:
            if a in cols:
                col_map[target] = a
                break

    year_score_cols = [c for c in df.columns if _extract_year_from_col(c) in {2023, 2024, 2025}]
    if not year_score_cols:
        return pd.DataFrame(
            columns=["year", "university", "department", "admission_type", "track_name", "percentile_type", "cutoff_score"]
        )

    rows: List[Dict] = []
    for _, r in df.iterrows():
        uni = str(r.get(col_map.get("university", ""), "")).strip()
        dept = str(r.get(col_map.get("department", ""), "")).strip()
        atype = str(r.get(col_map.get("admission_type", ""), "")).strip()
        track = str(r.get(col_map.get("track_name", ""), "")).strip()
        basis_text = str(r.get(col_map.get("basis", ""), ""))
        ptype = _detect_percentile_from_text(basis_text)
        for sc in year_score_cols:
            year = _extract_year_from_col(sc)
            score = pd.to_numeric(r.get(sc), errors="coerce")
            if pd.isna(score) or year is None:
                continue
            rows.append(
                {
                    "year": int(year),
                    "university": uni,
                    "department": dept,
                    "admission_type": atype,
                    "track_name": track,
                    "percentile_type": ptype,
                    "cutoff_score": float(score),
                }
            )
    out = pd.DataFrame(rows)
    if out.empty:
        return pd.DataFrame(
            columns=["year", "university", "department", "admission_type", "track_name", "percentile_type", "cutoff_score"]
        )
    for c in ["university", "department", "admission_type", "track_name"]:
        out[c] = _clean_text(out[c])
    return out


def load_data(dataset_mode: str = "merge") -> Tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    if not SUSI_2026.exists():
        pd.DataFrame(columns=["year", "university", "department", "admission_type", "track_name"]).to_csv(
            SUSI_2026, index=False, encoding="utf-8-sig"
        )
    if not SUSI_2027.exists():
        pd.DataFrame(columns=["year", "university", "department", "admission_type", "track_name"]).to_csv(
            SUSI_2027, index=False, encoding="utf-8-sig"
        )
    if not CUTOFFS.exists():
        pd.DataFrame(
            columns=["year", "university", "department", "admission_type", "track_name", "percentile_type", "cutoff_score"]
        ).to_csv(CUTOFFS, index=False, encoding="utf-8-sig")
    if not CRITERIA.exists():
        pd.DataFrame(columns=["university", "criterion", "weight", "description"]).to_csv(CRITERIA, index=False, encoding="utf-8-sig")

    s26 = _normalize_columns(load_csv(SUSI_2026), "susi")
    s27 = _normalize_columns(load_csv(SUSI_2027), "susi")
    s26["year"] = 2026
    s27["year"] = 2027

    for col in ["university", "department", "admission_type", "track_name"]:
        s26[col] = _clean_text(s26[col])
        s27[col] = _clean_text(s27[col])

    # 2027 대학명은 2026 기준 명칭으로 통일 (예: 강남대 -> 강남대학교)
    canon_map = _canonical_university_map(s26)
    s27["university"] = s27["university"].apply(
        lambda x: canon_map.get(_univ_key(x), str(x).strip())
    )

    if dataset_mode == "2027":
        merged = s27.copy()
    else:
        merged = pd.concat([s26, s27], ignore_index=True)
        key = ["university", "department", "admission_type", "track_name"]
        merged = merged.sort_values("year")
        merged = merged.drop_duplicates(subset=key, keep="last")  # 충돌 시 2027 우선
    merged = merged[merged["university"] != ""].copy()
    merged = merged[~merged.apply(lambda r: _remove_excluded_type(r["admission_type"], r["track_name"]), axis=1)].copy()

    cutoff = _normalize_cutoff_flexible(load_csv(CUTOFFS))
    for col in ["university", "department", "admission_type", "track_name"]:
        cutoff[col] = _clean_text(cutoff[col]) if col in cutoff.columns else ""
    cutoff["university"] = cutoff["university"].apply(lambda x: canon_map.get(_univ_key(x), str(x).strip()))
    cutoff = cutoff[cutoff["year"].between(2023, 2025)].copy()  # 2023~2025 기준

    criteria = load_csv(CRITERIA)
    return merged, cutoff, criteria


def extract_pdf_text(uploaded_file) -> str:
    if not uploaded_file:
        return ""
    reader = PdfReader(io.BytesIO(uploaded_file.getvalue()))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def extract_grade_points_from_pdf_text(pdf_text: str) -> List[float]:
    text = str(pdf_text or "")
    if not text.strip():
        return []
    pattern = re.compile(r"(?:(\d(?:\.\d+)?)\s*등급|등급\s*(\d(?:\.\d+)?))")
    vals: List[float] = []
    for m in pattern.finditer(text):
        a, b = m.groups()
        raw = a if a is not None else b
        try:
            v = float(raw)
        except Exception:
            continue
        if 1.0 <= v <= 9.0:
            vals.append(v)
    return vals


def extract_grade_points_from_excel(uploaded_file) -> List[float]:
    if not uploaded_file:
        return []
    vals: List[float] = []
    try:
        sheets = pd.read_excel(uploaded_file, sheet_name=None)
    except Exception:
        return []

    for _name, df in sheets.items():
        if df is None or df.empty:
            continue
        cols = [str(c) for c in df.columns]
        preferred = [c for c in cols if ("등급" in c) or ("내신" in c)]
        scan_cols = preferred if preferred else cols
        for c in scan_cols:
            series = pd.to_numeric(df[c], errors="coerce").dropna()
            if series.empty:
                continue
            for v in series.tolist():
                fv = float(v)
                if 1.0 <= fv <= 9.0:
                    vals.append(fv)
    return vals


def analyze_holistic_5level(pdf_text: str) -> Tuple[int, Dict[str, int], str]:
    text = (pdf_text or "").strip()
    if not text:
        detail = {
            "학업역량": 45,
            "전공적합성": 45,
            "자기주도성": 45,
            "공동체역량": 45,
            "발전가능성": 45,
        }
    else:
        base = min(100, len(text) // 60)
        keywords = {
            "학업역량": ["수업", "탐구", "학습", "과제", "교과"],
            "전공적합성": ["진로", "전공", "심화", "프로젝트", "연계"],
            "자기주도성": ["주도", "기획", "문제해결", "탐색", "개선"],
            "공동체역량": ["협업", "소통", "리더", "봉사", "배려"],
            "발전가능성": ["성장", "피드백", "도전", "확장", "변화"],
        }
        detail = {}
        for k, words in keywords.items():
            cnt = sum(text.count(w) for w in words)
            detail[k] = min(100, int(base * 0.5 + cnt * 9))

    avg = int(sum(detail.values()) / len(detail))
    if avg >= 85:
        level = 5
    elif avg >= 70:
        level = 4
    elif avg >= 55:
        level = 3
    elif avg >= 40:
        level = 2
    else:
        level = 1
    summary = f"학생부 분석 결과 5단계 중 {level}단계로 추정됩니다."
    return level, detail, summary


def get_cutoff_23_25(cutoffs: pd.DataFrame, uni: str, dept: str, atype: str, track: str) -> Tuple[float, float, str]:
    hit = cutoffs[
        (cutoffs["university"] == uni)
        & (cutoffs["department"] == dept)
        & (cutoffs["admission_type"] == atype)
        & (cutoffs["track_name"] == track)
    ]
    if hit.empty:
        return float("nan"), float("nan"), "none"

    c50 = hit.loc[hit["percentile_type"] == 50, "cutoff_score"].dropna()
    c70 = hit.loc[hit["percentile_type"] == 70, "cutoff_score"].dropna()

    cutoff50 = float(c50.median()) if not c50.empty else float("nan")
    cutoff70 = float(c70.median()) if not c70.empty else float("nan")
    if pd.notna(cutoff50):
        return cutoff50, cutoff70, "50"
    if pd.notna(cutoff70):
        return cutoff50, cutoff70, "70"
    return cutoff50, cutoff70, "none"


def rating_4level(student_grade: float, cutoff50: float, cutoff70: float, basis: str) -> Tuple[str, str]:
    # 내신 등급은 낮을수록 유리
    if basis == "none":
        return "입결 데이터 없음", "해당 조합의 2023~2025 50/70 컷 데이터가 없습니다."
    if basis == "50":
        if student_grade <= cutoff50:
            return "상향 가능", f"내신({student_grade:.2f})이 50% 컷({cutoff50:.2f}) 이내입니다."
        if pd.notna(cutoff70) and student_grade <= cutoff70:
            return "적정", f"내신({student_grade:.2f})이 70% 컷({cutoff70:.2f}) 이내입니다."
        if pd.notna(cutoff70) and student_grade <= cutoff70 + 0.5:
            return "소신", f"내신({student_grade:.2f})이 70% 컷({cutoff70:.2f})에 근접합니다."
        return "하향 권장", f"내신({student_grade:.2f})이 컷 대비 낮습니다."
    # basis == "70"
    if student_grade <= cutoff70:
        return "적정", f"50% 컷이 없어 70% 컷({cutoff70:.2f}) 기준으로 판단했습니다."
    if student_grade <= cutoff70 + 0.5:
        return "소신", f"50% 컷이 없어 70% 컷({cutoff70:.2f}) 기준으로 근접 판정했습니다."
    return "하향 권장", f"50% 컷이 없어 70% 컷({cutoff70:.2f}) 기준으로 하향 권장입니다."


def _basis_university(uni: str) -> str:
    return uni if uni in TOP15 else "경희대"


def build_report_text(payload: Dict, choices: List[SupportChoice], holistic_detail: Dict[str, int]) -> str:
    lines = [
        "# 나의 입시 위치 진단 종합보고서",
        f"- 생성시각: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        "",
        "## 1) 사용자 정보",
        f"- 컨설턴트: {payload.get('consultant_name', '')}",
        f"- 학생명: {payload.get('student_name', '')}",
        f"- 학교명: {payload.get('school_name', '')}",
        f"- 학년: {payload.get('grade', '')}",
        f"- 학생 연락처: {payload.get('student_phone', '')}",
        f"- 이메일: {payload.get('email', '')}",
        f"- 학부모 연락처: {payload.get('parent_phone', '')}",
        f"- 내신 등급(평균): {payload.get('student_grade_score', '')}",
        f"- 내신 산출 기준: {payload.get('student_grade_source', '')}",
        "",
        "## 2) 학생부 분석 (5단계)",
        f"- 종합단계: {payload.get('holistic_level', 0)}단계",
    ]
    for k, v in holistic_detail.items():
        lines.append(f"- {k}: {v}점")

    lines.extend(["", "## 3) 지원희망대학 평가 (최대 6개)"])
    for c in choices:
        basis = _basis_university(c.university)
        lines.append(
            f"- {c.support_no}. {c.university} / {c.department} / {c.admission_type} / {c.track_name}"
            f" -> {c.diag_level} (50컷 {c.cutoff50 if pd.notna(c.cutoff50) else '없음'}, "
            f"70컷 {c.cutoff70 if pd.notna(c.cutoff70) else '없음'}, "
            f"판단기준 {c.cutoff_basis}, 기준대학 {basis})"
        )
        lines.append(f"  - 판단근거: {c.diag_reason}")

    lines.extend(
        [
            "",
            "## 4) 적용 기준",
            "- 대학 리스트: 2026 수시검색기 + 2027 업로드 병합",
            "- 충돌 규칙: 동일 조합(대학교/모집단위/전형유형/전형명) 충돌 시 2027 우선",
            "- 입결 판단: 2023~2025학년도 50%/70% 컷 기준",
            "- 학생부종합 평가기준: 서울 상위 15개 대학 기준, 미포함 대학은 경희대 기준 적용",
        ]
    )
    return "\n".join(lines)


def build_docx_bytes(payload: Dict, choices: List[SupportChoice], holistic_detail: Dict[str, int]) -> bytes:
    doc = Document()
    doc.add_heading("나의 입시 위치 진단 종합보고서", level=1)
    doc.add_paragraph(f"생성시각: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    doc.add_heading("1) 사용자 정보", level=2)
    info_rows = [
        ("컨설턴트", payload.get("consultant_name", "")),
        ("학생명", payload.get("student_name", "")),
        ("학교명", payload.get("school_name", "")),
        ("학년", payload.get("grade", "")),
        ("학생 연락처", payload.get("student_phone", "")),
        ("이메일", payload.get("email", "")),
        ("학부모 연락처", payload.get("parent_phone", "")),
        ("내신 등급(평균)", payload.get("student_grade_score", "")),
        ("내신 산출 기준", payload.get("student_grade_source", "")),
    ]
    for k, v in info_rows:
        doc.add_paragraph(f"- {k}: {v}")

    doc.add_heading("2) 학생부 분석 (5단계)", level=2)
    doc.add_paragraph(f"- 종합단계: {payload.get('holistic_level', 0)}단계")
    for k, v in holistic_detail.items():
        doc.add_paragraph(f"- {k}: {v}점")

    doc.add_heading("3) 지원희망대학 평가", level=2)
    table = doc.add_table(rows=1, cols=8)
    hdr = table.rows[0].cells
    hdr[0].text = "번호"
    hdr[1].text = "대학/모집단위"
    hdr[2].text = "전형"
    hdr[3].text = "평가"
    hdr[4].text = "50컷"
    hdr[5].text = "70컷"
    hdr[6].text = "판단근거"
    hdr[7].text = "판단기준"

    for c in choices:
        row = table.add_row().cells
        row[0].text = str(c.support_no)
        row[1].text = f"{c.university} / {c.department}"
        row[2].text = f"{c.admission_type} / {c.track_name}"
        row[3].text = c.diag_level
        row[4].text = f"{c.cutoff50:.2f}" if pd.notna(c.cutoff50) else "없음"
        row[5].text = f"{c.cutoff70:.2f}" if pd.notna(c.cutoff70) else "없음"
        row[6].text = c.diag_reason
        row[7].text = c.cutoff_basis

    doc.add_heading("4) 적용 기준", level=2)
    doc.add_paragraph("- 대학 리스트: 2026 수시검색기 + 2027 업로드 병합 (충돌 시 2027 우선)")
    doc.add_paragraph("- 입결 판단: 2023~2025학년도 50%/70% 컷 기준")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def save_session(payload: Dict, choices: List[SupportChoice]) -> None:
    conn = sqlite3.connect(DB_PATH)
    try:
        cur = conn.cursor()
        outcome = {"payload": payload, "supports": [c.__dict__ for c in choices]}
        cur.execute(
            """
            INSERT INTO diagnosis_sessions (
                created_at, consultant_name, student_name, school_name, grade,
                student_phone, email, parent_phone, student_index, pdf_summary,
                holistic_score, outcome_json
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                datetime.now().isoformat(timespec="seconds"),
                payload.get("consultant_name", ""),
                payload.get("student_name", ""),
                payload.get("school_name", ""),
                payload.get("grade", ""),
                payload.get("student_phone", ""),
                payload.get("email", ""),
                payload.get("parent_phone", ""),
                float(payload.get("student_index", 0)),
                payload.get("pdf_summary", ""),
                float(payload.get("holistic_score", 0)),
                json.dumps(outcome, ensure_ascii=False),
            ),
        )
        session_id = int(cur.lastrowid)
        for c in choices:
            cur.execute(
                """
                INSERT INTO support_choices (
                    session_id, support_no, university, department, admission_type,
                    track_name, diag_level, diag_reason, cutoff50, cutoff70
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    session_id,
                    c.support_no,
                    c.university,
                    c.department,
                    c.admission_type,
                    c.track_name,
                    c.diag_level,
                    c.diag_reason,
                    c.cutoff50,
                    c.cutoff70,
                ),
            )
        conn.commit()
    finally:
        conn.close()


def _step_header(current_step: int) -> None:
    labels = ["1. 사용자정보", "2. 학생부 분석", "3. 희망대학 지원", "4. 종합 보고서"]
    cols = st.columns(4)
    for i, (col, label) in enumerate(zip(cols, labels), start=1):
        with col:
            cls = "step-box step-active" if i == current_step else "step-box"
            st.markdown(f"<div class='{cls}'>{label}</div>", unsafe_allow_html=True)


def _consultant_panel() -> str:
    if "consultant_name_fixed" not in st.session_state:
        st.session_state.consultant_name_fixed = ""

    if st.session_state.consultant_name_fixed:
        st.text_input("컨설턴트 이름", value=st.session_state.consultant_name_fixed, disabled=True, key="consultant_fixed_view")
        return st.session_state.consultant_name_fixed

    c1, c2 = st.columns([3, 1])
    with c1:
        name = st.text_input("컨설턴트 이름", key="consultant_name_once")
    with c2:
        save = st.button("저장", use_container_width=True)
    if save and name.strip():
        st.session_state.consultant_name_fixed = name.strip()
        st.rerun()
    return name.strip()


def _choice_input_block(susi_df: pd.DataFrame, no: int) -> Tuple[str, str, str, str]:
    st.markdown(f"지원 {no}")
    a, b, c, d = st.columns(4)
    uni_options = sorted(susi_df["university"].dropna().unique().tolist())

    with a:
        uni = st.selectbox(f"대학교 {no}", ["직접입력"] + uni_options, key=f"uni_{no}")
        uni_text = st.text_input(f"대학교 직접입력 {no}", key=f"uni_txt_{no}") if uni == "직접입력" else uni

    filtered_uni = susi_df[susi_df["university"] == uni_text] if uni_text else susi_df
    dept_options = sorted(filtered_uni["department"].dropna().unique().tolist()) or [""]

    with b:
        dept = st.selectbox(f"모집단위 {no}", ["직접입력"] + dept_options, key=f"dept_{no}")
        dept_text = st.text_input(f"모집단위 직접입력 {no}", key=f"dept_txt_{no}") if dept == "직접입력" else dept

    filtered_dept = filtered_uni[filtered_uni["department"] == dept_text] if dept_text else filtered_uni
    atype_options = sorted(filtered_dept["admission_type"].dropna().unique().tolist())
    track_options = sorted(filtered_dept["track_name"].dropna().unique().tolist()) or [""]

    with c:
        atype = st.selectbox(f"전형유형 {no}", atype_options or ADMISSION_TYPES, key=f"atype_{no}")
    with d:
        track = st.selectbox(f"전형명 {no}", ["직접입력"] + track_options, key=f"track_{no}")
        track_text = st.text_input(f"전형명 직접입력 {no}", key=f"track_txt_{no}") if track == "직접입력" else track

    return uni_text.strip(), dept_text.strip(), atype.strip(), track_text.strip()


def main() -> None:
    st.set_page_config(page_title="나의 입시 위치 진단 서비스", layout="wide")
    init_db()
    inject_css()

    st.markdown("<div class='service-title'>나의 입시 위치 진단서비스</div>", unsafe_allow_html=True)
    st.markdown("<div class='service-subtitle'>by 대치수프리마</div>", unsafe_allow_html=True)
    meta_l, meta_r = st.columns([3, 4])
    with meta_l:
        st.markdown(
            "<div class='top-meta-left'>기준: 2026+2027 수시 병합(충돌 시 2027 우선), 2023~2025 50/70컷</div>",
            unsafe_allow_html=True,
        )
    with meta_r:
        st.markdown(
            "<div class='top-meta-right'>대학기준: 2026 183개 대학 및 2027 수시 모집요강 발표한 서울경기지역 학교</div>",
            unsafe_allow_html=True,
        )

    with st.sidebar:
        st.markdown("### 운영 설정")
        dataset_mode_label = st.radio(
            "대학 데이터 사용 방식",
            ["2026+2027 병합(충돌 시 2027 우선)", "2027만 사용(전체 전환)"],
            index=0,
            help="2027학년도 전체 데이터 발표 후에는 '2027만 사용'으로 운영 가능합니다.",
        )
        dataset_mode = "2027" if dataset_mode_label.startswith("2027만") else "merge"

        with st.expander("데이터 관리 (필요 시 열람/교체)", expanded=False):
            c1, c2 = st.columns(2)
            with c1:
                st.download_button("2026 수시 템플릿", SUSI_2026.read_bytes() if SUSI_2026.exists() else b"", file_name="susi_explorer.csv")
                up1 = st.file_uploader("2026 수시 CSV 업로드", type=["csv"], key="up_s26")
                if up1:
                    save_uploaded_csv(up1, SUSI_2026)
                    st.success("2026 수시 데이터 반영 완료")
            with c2:
                st.download_button("2027 수시 템플릿", SUSI_2027.read_bytes() if SUSI_2027.exists() else b"", file_name="susi_explorer_2027.csv")
                up2 = st.file_uploader("2027 수시 CSV 업로드", type=["csv"], key="up_s27")
                if up2:
                    save_uploaded_csv(up2, SUSI_2027)
                    st.success("2027 수시 데이터 반영 완료")
            c3, c4 = st.columns(2)
            with c3:
                st.download_button("컷 템플릿", CUTOFFS.read_bytes() if CUTOFFS.exists() else b"", file_name="admission_cutoffs.csv")
                up3 = st.file_uploader("컷 CSV 업로드(2023~2025)", type=["csv"], key="up_cut")
                if up3:
                    save_uploaded_csv(up3, CUTOFFS)
                    st.success("컷 데이터 반영 완료")
            with c4:
                st.download_button("학생부 평가기준 템플릿", CRITERIA.read_bytes() if CRITERIA.exists() else b"", file_name="holistic_criteria.csv")
                up4 = st.file_uploader("평가기준 CSV 업로드", type=["csv"], key="up_criteria")
                if up4:
                    save_uploaded_csv(up4, CRITERIA)
                    st.success("평가기준 데이터 반영 완료")

    susi_df, cutoff_df, _criteria_df = load_data(dataset_mode=dataset_mode)

    if "step" not in st.session_state:
        st.session_state.step = 1
    if "profile" not in st.session_state:
        st.session_state.profile = {}
    if "holistic" not in st.session_state:
        st.session_state.holistic = {}
    if "supports" not in st.session_state:
        st.session_state.supports = []
    if "saved" not in st.session_state:
        st.session_state.saved = False

    head_left, head_right = st.columns([6, 2])
    with head_left:
        _step_header(st.session_state.step)
    with head_right:
        consultant_global = _consultant_panel()
    st.markdown("---")

    # Step 1
    if st.session_state.step == 1:
        st.subheader("1단계 - 사용자정보")
        with st.form("step1_form"):
            c1, c2 = st.columns(2)
            with c1:
                student_name = st.text_input("학생명 *")
                school_name = st.text_input("학교명 *")
                grade = st.selectbox("학년 *", GRADE_OPTIONS)
            with c2:
                student_phone = st.text_input("학생 전화번호 *")
                email = st.text_input("메일주소(필수) *")
                parent_phone = st.text_input("학부모 연락처 *")

            _, btn_col = st.columns([5, 1])
            with btn_col:
                ok = st.form_submit_button("다음 단계", use_container_width=True)

        if ok:
            required = {
                "컨설턴트 이름": consultant_global,
                "학생명": student_name,
                "학교명": school_name,
                "학생 전화번호": student_phone,
                "메일주소": email,
                "학부모 연락처": parent_phone,
            }
            missing = [k for k, v in required.items() if not str(v).strip()]
            if missing:
                st.error("필수값 누락: " + ", ".join(missing))
            else:
                st.session_state.profile = {
                    "consultant_name": consultant_global.strip(),
                    "student_name": student_name.strip(),
                    "school_name": school_name.strip(),
                    "grade": grade.strip(),
                    "student_phone": student_phone.strip(),
                    "email": email.strip(),
                    "parent_phone": parent_phone.strip(),
                }
                st.session_state.step = 2
                st.rerun()

    # Step 2
    elif st.session_state.step == 2:
        st.subheader("2단계 - 학생부 분석")
        lpad, center, rpad = st.columns([1, 2, 1])
        with center:
            pdf_file = st.file_uploader("학생부 PDF 업로드", type=["pdf"], key="pdf_step2")
        excel_file = st.file_uploader("내신 계산 엑셀 업로드(xlsx)", type=["xlsx"], key="grade_xlsx_step2")
        g1, g2 = st.columns(2)
        with g1:
            grades_text = st.text_input("과목별 내신 등급(쉼표 구분)", placeholder="예: 2.1, 2.3, 1.9, 2.4")
        with g2:
            manual_grade = st.number_input("내신 평균(직접 입력, 선택)", min_value=1.0, max_value=9.0, value=2.5, step=0.01)

        calc_grade = None
        calc_source = ""
        if grades_text.strip():
            try:
                vals = [float(x.strip()) for x in grades_text.split(",") if x.strip()]
                if vals:
                    calc_grade = round(sum(vals) / len(vals), 2)
                    calc_source = f"수기 입력({len(vals)}개)"
                    st.caption(f"자동 계산 내신 평균: {calc_grade}")
            except Exception:
                st.warning("내신 등급 형식이 올바르지 않습니다. 예: 2.1, 2.3, 1.9")

        c1, c2, c3 = st.columns(3)
        with c1:
            if st.button("이전 단계"):
                st.session_state.step = 1
                st.rerun()
        with c2:
            if st.button("분석 실행", use_container_width=True):
                text = extract_pdf_text(pdf_file) if pdf_file else ""
                level, detail, summary = analyze_holistic_5level(text)
                excel_points = extract_grade_points_from_excel(excel_file)
                pdf_points = extract_grade_points_from_pdf_text(text)

                if excel_points:
                    grade_score = round(sum(excel_points) / len(excel_points), 2)
                    grade_source = f"엑셀 계산({len(excel_points)}개)"
                elif calc_grade is not None:
                    grade_score = calc_grade
                    grade_source = calc_source or "수기 입력 계산"
                elif pdf_points:
                    grade_score = round(sum(pdf_points) / len(pdf_points), 2)
                    grade_source = f"학생부 PDF 추출({len(pdf_points)}개)"
                else:
                    grade_score = float(manual_grade)
                    grade_source = "직접 입력"

                st.session_state.holistic = {
                    "pdf_summary": summary,
                    "holistic_level": level,
                    "holistic_detail": detail,
                    "holistic_score": sum(detail.values()) / len(detail),
                    "student_grade_score": grade_score,
                    "student_grade_source": grade_source,
                }
                st.success("학생부 분석이 완료되었습니다. 아래에서 결과를 확인하세요.")
        with c3:
            if st.button("다음 단계", use_container_width=True):
                if not st.session_state.holistic:
                    st.warning("먼저 분석 실행을 눌러 주세요.")
                else:
                    st.session_state.step = 3
                    st.rerun()

        if st.session_state.holistic:
            st.markdown("### 분석 결과")
            st.info(st.session_state.holistic.get("pdf_summary", ""))
            detail = st.session_state.holistic.get("holistic_detail", {})
            m1, m2 = st.columns(2)
            with m1:
                st.metric("종합 단계", f"{st.session_state.holistic.get('holistic_level', 0)}단계")
            with m2:
                st.metric("평균 점수", f"{st.session_state.holistic.get('holistic_score', 0):.1f}")
            st.metric("내신 평균", f"{st.session_state.holistic.get('student_grade_score', 0):.2f}")
            st.caption(f"내신 산출 기준: {st.session_state.holistic.get('student_grade_source', '-')}")
            for k, v in detail.items():
                st.write(f"- {k}: {v}점")
                st.progress(min(max(int(v), 0), 100))
        else:
            st.markdown("### 분석 결과")
            st.info("학생부 PDF를 업로드하고 '분석 실행'을 누르면 이 영역에 결과가 표시됩니다.")

    # Step 3
    elif st.session_state.step == 3:
        st.subheader("3단계 - 희망대학 지원 (최대 6개)")
        supports: List[Tuple[str, str, str, str]] = []

        with st.form("step3_form"):
            for no in range(1, 7):
                with st.expander(f"지원 {no}", expanded=(no <= 2)):
                    supports.append(_choice_input_block(susi_df, no))

            c1, c2 = st.columns(2)
            with c1:
                back = st.form_submit_button("이전 단계")
            with c2:
                nxt = st.form_submit_button("평가 후 보고서 보기", use_container_width=True)

        if back:
            st.session_state.step = 2
            st.rerun()

        if nxt:
            rows: List[SupportChoice] = []
            student_grade = float(st.session_state.holistic.get("student_grade_score", 0) or 0)
            for i, (uni, dept, atype, track) in enumerate(supports, start=1):
                if not (uni and dept and atype and track):
                    continue
                c50, c70, basis = get_cutoff_23_25(cutoff_df, uni, dept, atype, track)
                level, reason = rating_4level(student_grade, c50, c70, basis)
                rows.append(SupportChoice(i, uni, dept, atype, track, level, reason, c50, c70, basis))

            if not rows:
                st.warning("최소 1개 이상의 지원 대학/학과/전형 정보를 입력해 주세요.")
            else:
                st.session_state.supports = [r.__dict__ for r in rows]
                st.session_state.step = 4
                st.rerun()

    # Step 4
    else:
        st.subheader("4단계 - 종합 보고서")

        payload = {
            **st.session_state.profile,
            **st.session_state.holistic,
            "student_index": st.session_state.holistic.get("student_grade_score", 0),
            "top15_reference": TOP15,
        }
        choices = [SupportChoice(**x) for x in st.session_state.supports]
        holistic_detail = st.session_state.holistic.get("holistic_detail", {})

        if not st.session_state.saved:
            save_session(payload, choices)
            st.session_state.saved = True

        report_text = build_report_text(payload, choices, holistic_detail)
        st.text_area("종합 보고서 미리보기", value=report_text, height=420)

        docx_bytes = build_docx_bytes(payload, choices, holistic_detail)
        stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        st.download_button(
            "DOC 보고서 다운로드",
            data=docx_bytes,
            file_name=f"admission_report_{stamp}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        c1, c2 = st.columns(2)
        with c1:
            if st.button("이전 단계"):
                st.session_state.step = 3
                st.rerun()
        with c2:
            if st.button("새 진단 시작", use_container_width=True):
                st.session_state.step = 1
                st.session_state.profile = {}
                st.session_state.holistic = {}
                st.session_state.supports = []
                st.session_state.saved = False
                st.rerun()

        st.info(f"진단 결과가 DB에 저장되었습니다: {DB_PATH}")

    st.markdown("<div class='footer-note'>자료: 대학어디가  운영: 대치 수프리마</div>", unsafe_allow_html=True)


if __name__ == "__main__":
    main()


